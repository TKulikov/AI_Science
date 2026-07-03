"""
Загрузка и парсинг документов.
Метаданные извлекаются из структуры пути (журнал, год, тип и т.п.)
"""

from typing import Iterator
from dataclasses import dataclass
from pathlib import Path
from typing import Optional
import re


SOURCE_TYPE_MAP = {
    "журналы": "journal",
    "доклады": "report",
    "статьи": "article",
    "обзоры": "review",
    "материалы конференций": "conference",
    "патенты": "patent",
}

KNOWN_JOURNALS = {
    "горная промышленность",
    "горный журнал",
    "обогащение руд",
    "цветные металлы",
}


# класс документа
@dataclass
class RawDocument:
    path: Path
    text: str
    language: str # ru | en | mixed
    source_type: str # journal | report ...
    journal: Optional[str] # "цветные металлы" | None ...
    year: Optional[int] # 2024 | None ...
    geography: str # "RU" по умолчанию
    file_format: str # "pdf" | "docx" | "txt"
    

# метаданные из структуры пути

def parse_year(part: str) -> Optional[int]:
    # извлекаеем год типа '2024'
    m = re.match(r"(\d{4})", part)
    return int(m.group(1)) if m else None


def detect_geography(parts: list[str]) -> str:
    # определяем географию по ключевым словам
    ru_markers = {"горная промышленность", "горный журнал", "обогащение руд",
                  "цветные металлы", "доклады", "обзоры", "статьи"}
    lowered = {p.lower() for p in parts}
    return "RU" if lowered & ru_markers else "foreign"


def extract_path_metadata(path: Path, root: Path):
    # извлекаем метаданные документа из его относительного пути
    
    rel_parts = [p.lower() for p in path.relative_to(root).parts[:-1]]
    
    source_type = "unknown"
    for part in rel_parts:
        for k, v in SOURCE_TYPE_MAP.items():
            if k in part:
                source_type = v
                break
            
    journal = None
    for part in rel_parts:
        if part in KNOWN_JOURNALS:
            journal = part
            break
        
    year = None
    for part in rel_parts:
        y = parse_year(part)
        if y:
            year = y
            break
        
    geography = detect_geography(rel_parts)
    
    return {
        "source_type": source_type,
        "journal": journal,
        "year": year,
        "geography": geography,
    }
    

# определение языка
CYRILLIC = re.compile(r"[а-яёА-ЯЁ]")
LATIN = re.compile(r"a-zA-Z")

def detect_language(text: str) -> str:
    # определяем язык текста по соотношению кириллицы и латиницы
    
    cyr = len(CYRILLIC.findall(text))
    lat = len(LATIN.findall(text))
    
    total = cyr + lat
    if total == 0:
        return "ru"
    
    rat = cyr / total
    # 70%
    if rat > 0.7:
        return "ru"
    if rat < 0.3:
        return "en"
    
    # иначе непонятно
    return "mixed"


# парсеры 

def parse_pdf(path: Path) -> str:
    
    # некоторые pdf файлы требуют cryptography>=3.1
    
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        pages = (page.extract_text() or "" for page in reader.pages)
        return "\n".join(pages)
    except Exception as e:
        raise RuntimeError(f"Не удалось прочитать PDF {path}: {e}")
    
def parse_docx(path: Path) -> str:
    
    # TODO все doc файлы конвертировать в docx, чтобы убрать ошибку:
    # .doc is not a Word file, content type is 'application/vnd.openxmlformats-officedocument.themeManager+xml'
    
    try:
        from docx import Document
        doc = Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as e:
        raise RuntimeError(f"Не удалось прочитать DOCX {path}: {e}")
    
def parse_txt(path: Path) -> str:
    for i in ("utf-8", "cp1251"):
        try:
            return path.read_text(encoding=i)
        except UnicodeDecodeError:
            continue
    raise RuntimeError(f"Не удалось прочитать текстовый файл {path}")


# вроде бы все статьи в корпусе именно этого формата (неуверен насчёт txt и md, но пусть будет)

PARSERS = {
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".doc": parse_docx,
    ".txt": parse_txt,
    ".md": parse_txt,
}


def parse_file(path: Path) -> str:
    ext = path.suffix.lower()
    parser = PARSERS.get(ext)
    if parser is None:
        raise ValueError(f"Неподдерживаемый формат: {ext} ({path})")
    return parser(path)


#-------------------
# грузим документы
#-------------------

def load_document(path: Path, root: Path) -> RawDocument:
    # загружает один документ и возвращает RawDocument с метаданными
    
    text = parse_file(path)
    meta = extract_path_metadata(path, root)
    lang = detect_language(text)
    
    return RawDocument(
        path=path,
        text=text,
        language=lang,
        source_type=meta["source_type"],
        journal=meta["journal"],
        year=meta["year"],
        geography=meta["geography"],
        file_format=path.suffix.lower().lstrip("."),
    )
    

def iter_documents(root: Path) -> Iterator[RawDocument]:
    # рекурсивно обходит директорию и загружает все поддерживаемые документы
    
    supported = set(PARSERS.keys())
    
    for path in sorted(root.rglob("*")):
        if not path.is_file():
            continue
        if path.suffix.lower() not in supported:
            continue
        
        try:
            yield load_document(path, root)
        except Exception as e:
            print(f"Пропущен {path.name}: {e}")